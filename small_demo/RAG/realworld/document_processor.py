"""
文档处理模块 (Document Processing Module)

这个模块负责处理各种格式的文档，包括文本文件、PDF、Word 文档等。
提供统一的文档加载、解析和分块接口。
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from abc import ABC, abstractmethod
import logging

from config import get_config
import logger

class Document:
    """文档类，封装文档的基本信息和内容"""

    def __init__(self, content: str, metadata: Optional[Dict[str, Any]] = None, source: Optional[str] = None):
        """
        初始化文档

        参数:
            content: 文档内容
            metadata: 元数据字典
            source: 文档来源路径
        """
        self.content = content
        self.metadata = metadata or {}
        self.source = source

        # 自动添加一些元数据
        if source:
            self.metadata['source'] = source
            self.metadata['file_size'] = os.path.getsize(source) if os.path.exists(source) else 0

    def __str__(self) -> str:
        return f"Document(source={self.source}, content_length={len(self.content)})"

class DocumentProcessor(ABC):
    """文档处理器抽象基类"""

    @abstractmethod
    def can_process(self, file_path: str) -> bool:
        """检查是否能处理指定文件"""
        pass

    @abstractmethod
    def process(self, file_path: str) -> Document:
        """处理文档并返回 Document 对象"""
        pass

class TextDocumentProcessor(DocumentProcessor):
    """文本文档处理器"""

    def can_process(self, file_path: str) -> bool:
        """检查文件扩展名"""
        return Path(file_path).suffix.lower() in ['.txt', '.md']

    def process(self, file_path: str) -> Document:
        """处理文本文件"""
        config = get_config()
        logger = logging.getLogger(__name__)

        try:
            with open(file_path, 'r', encoding=config.document.encoding, errors='ignore') as f:
                content = f.read()

            # 提取元数据
            metadata = {
                'file_type': Path(file_path).suffix,
                'encoding': config.document.encoding,
                'line_count': len(content.split('\n'))
            }

            logger.info(f"成功处理文本文档: {file_path}")
            return Document(content=content, metadata=metadata, source=file_path)

        except Exception as e:
            logger.error(f"处理文本文档失败 {file_path}: {e}")
            raise

class PDFDocumentProcessor(DocumentProcessor):
    """PDF 文档处理器"""

    def __init__(self):
        try:
            import PyPDF2
            self.pdf_reader = PyPDF2
        except ImportError:
            raise ImportError("处理 PDF 需要安装 PyPDF2: pip install PyPDF2")

    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.pdf'

    def process(self, file_path: str) -> Document:
        """处理 PDF 文件"""
        logger = logging.getLogger(__name__)

        try:
            with open(file_path, 'rb') as f:
                pdf_reader = self.pdf_reader.PdfReader(f)
                content = ""

                # 提取所有页面的文本
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    content += page.extract_text() + "\n"

            # 提取元数据
            metadata = {
                'file_type': '.pdf',
                'page_count': len(pdf_reader.pages),
                'title': pdf_reader.metadata.get('/Title', ''),
                'author': pdf_reader.metadata.get('/Author', ''),
                'subject': pdf_reader.metadata.get('/Subject', '')
            }

            logger.info(f"成功处理 PDF 文档: {file_path} ({len(pdf_reader.pages)} 页)")
            return Document(content=content, metadata=metadata, source=file_path)

        except Exception as e:
            logger.error(f"处理 PDF 文档失败 {file_path}: {e}")
            raise

class WordDocumentProcessor(DocumentProcessor):
    """Word 文档处理器"""

    def __init__(self):
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise ImportError("处理 Word 文档需要安装 python-docx: pip install python-docx")

    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in ['.docx', '.doc']

    def process(self, file_path: str) -> Document:
        """处理 Word 文档"""
        logger = logging.getLogger(__name__)

        try:
            doc = self.docx.Document(file_path)
            content = ""

            # 提取所有段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"

            # 提取元数据
            metadata = {
                'file_type': Path(file_path).suffix,
                'paragraph_count': len(doc.paragraphs),
                'word_count': len(content.split())
            }

            logger.info(f"成功处理 Word 文档: {file_path}")
            return Document(content=content, metadata=metadata, source=file_path)

        except Exception as e:
            logger.error(f"处理 Word 文档失败 {file_path}: {e}")
            raise

class TextSplitter:
    """文本分割器"""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        初始化文本分割器

        参数:
            chunk_size: 每个块的最大字符数
            chunk_overlap: 相邻块之间的重叠字符数
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.logger = logging.getLogger(__name__)

    def split_text(self, text: str) -> List[str]:
        """
        将文本分割成块

        参数:
            text: 要分割的文本

        返回:
            分割后的文本块列表
        """
        if not text:
            return []

        chunks = []
        start = 0

        while start < len(text):
            # 计算块的结束位置
            end = start + self.chunk_size

            # 如果不是最后一块，确保在句子边界结束
            if end < len(text):
                # 寻找最近的句子结束符
                sentence_endings = ['. ', '! ', '? ', '\n\n']
                best_end = end

                for ending in sentence_endings:
                    last_ending = text.rfind(ending, start, end)
                    if last_ending > start and abs(last_ending - end) < abs(best_end - end):
                        best_end = last_ending + len(ending)

                end = best_end

            # 提取块
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # 计算下一个块的起始位置（考虑重叠）
            start = max(start + 1, end - self.chunk_overlap)

        self.logger.debug(f"文本分割完成: {len(chunks)} 个块")
        return chunks

    def split_document(self, document: Document) -> List[Document]:
        """
        分割文档为多个块

        参数:
            document: 要分割的文档

        返回:
            分割后的文档块列表
        """
        chunks = self.split_text(document.content)

        result = []
        for i, chunk in enumerate(chunks):
            # 创建新的元数据
            chunk_metadata = document.metadata.copy()
            chunk_metadata.update({
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            })

            result.append(Document(
                content=chunk,
                metadata=chunk_metadata,
                source=document.source
            ))

        self.logger.info(f"文档分割完成: {document.source} -> {len(result)} 个块")
        return result

class DocumentLoader:
    """文档加载器"""

    def __init__(self):
        """初始化文档处理器列表"""
        self.processors = [
            TextDocumentProcessor(),
            PDFDocumentProcessor(),
            WordDocumentProcessor()
        ]
        self.logger = logging.getLogger(__name__)

    def load_document(self, file_path: str) -> Document:
        """
        加载单个文档

        参数:
            file_path: 文档文件路径

        返回:
            加载的文档对象

        异常:
            如果文件不存在或不支持的格式，抛出异常
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 查找合适的处理器
        for processor in self.processors:
            if processor.can_process(file_path):
                self.logger.info(f"使用 {processor.__class__.__name__} 处理文件: {file_path}")
                return processor.process(file_path)

        # 没有找到合适的处理器
        supported_extensions = []
        for processor in self.processors:
            # 这里需要一个方法来获取支持的扩展名，暂时简化
            pass

        raise ValueError(f"不支持的文件格式: {file_path}")

    def load_documents(self, directory: str, recursive: bool = True) -> List[Document]:
        """
        加载目录中的所有文档

        参数:
            directory: 文档目录路径
            recursive: 是否递归加载子目录

        返回:
            加载的文档列表
        """
        config = get_config()
        dir_path = Path(directory)

        if not dir_path.exists():
            raise FileNotFoundError(f"目录不存在: {directory}")

        documents = []

        # 构建文件模式
        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"

        # 遍历文件
        for file_path in dir_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in config.document.supported_extensions:
                try:
                    doc = self.load_document(str(file_path))
                    documents.append(doc)
                    self.logger.info(f"加载文档: {file_path}")
                except Exception as e:
                    self.logger.warning(f"跳过文件 {file_path}: {e}")

        self.logger.info(f"从目录 {directory} 加载了 {len(documents)} 个文档")
        return documents

def create_text_splitter() -> TextSplitter:
    """创建配置的文本分割器"""
    config = get_config()
    return TextSplitter(
        chunk_size=config.document.chunk_size,
        chunk_overlap=config.document.chunk_overlap
    )

def create_document_loader() -> DocumentLoader:
    """创建文档加载器"""
    return DocumentLoader()